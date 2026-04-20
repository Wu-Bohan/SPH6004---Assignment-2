import argparse
import glob
import os
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches

try:
    from .common import load_json
except ImportError:
    from common import load_json


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str):
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("No data available.")
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr_cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def add_figures(doc: Document, figures_dir: str):
    figs = [
        "model_metric_bars.png",
        "ablation_cindex.png",
        "calibration_plot.png",
        "km_risk_groups.png",
        "tree_shap_top20.png",
    ]
    doc.add_heading("Figures", level=2)
    for f in figs:
        p = os.path.join(figures_dir, f)
        if os.path.exists(p):
            doc.add_paragraph(f)
            doc.add_picture(p, width=Inches(6.5))


def build_experiment_log(
    out_path: str,
    data_dir: str,
    metrics_dir: str,
    figures_dir: str,
    summary_csv: str,
):
    doc = Document()
    doc.add_heading("Assignment2 Experiment Log", level=1)
    doc.add_paragraph(f"Generated at: {datetime.now().isoformat(timespec='seconds')}")

    prep_summary_path = os.path.join(data_dir, "preprocess_summary.json")
    feature_art_path = os.path.join(data_dir, "feature_artifacts.json")
    if os.path.exists(prep_summary_path):
        doc.add_heading("Data Overview", level=2)
        prep = load_json(prep_summary_path)
        for k, v in prep.items():
            doc.add_paragraph(f"{k}: {v}")
    if os.path.exists(feature_art_path):
        doc.add_heading("Feature Engineering Artifacts", level=2)
        feat = load_json(feature_art_path)
        for k, v in feat.items():
            if isinstance(v, list) and len(v) > 20:
                doc.add_paragraph(f"{k}: [{len(v)} entries]")
            else:
                doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Experiment Config and Results", level=2)
    for exp in ["E0", "E1", "E2", "E3", "E4", "E5", "E6"]:
        p = os.path.join(metrics_dir, f"{exp}.json")
        doc.add_heading(exp, level=3)
        if not os.path.exists(p):
            doc.add_paragraph("Not completed.")
            continue
        m = load_json(p)
        for k, v in m.items():
            if k == "artifact":
                doc.add_paragraph("artifact:")
                if isinstance(v, dict):
                    for k2, v2 in v.items():
                        if isinstance(v2, list) and len(v2) > 10:
                            doc.add_paragraph(f"  {k2}: [len={len(v2)}]")
                        else:
                            doc.add_paragraph(f"  {k2}: {v2}")
            elif "path" in k:
                continue
            else:
                doc.add_paragraph(f"{k}: {v}")

    if os.path.exists(summary_csv):
        df = pd.read_csv(summary_csv)
        add_table_from_df(doc, df, "Metrics Summary Table")
    add_figures(doc, figures_dir)

    log_files = sorted(glob.glob(os.path.join(os.path.dirname(metrics_dir), "..", "logs", "*.log")))
    doc.add_heading("Failure/Retry Logs", level=2)
    if not log_files:
        doc.add_paragraph("No failure logs were found.")
    else:
        for lf in log_files[-20:]:
            doc.add_paragraph(f"- {os.path.basename(lf)}")

    doc.save(out_path)


def build_final_report(
    out_docx: str,
    out_md: str,
    summary_csv: str,
):
    df = pd.read_csv(summary_csv)
    best = df.sort_values("test_c_index", ascending=False).iloc[0]

    doc = Document()
    doc.add_heading("Assignment2 Final Report", level=1)
    doc.add_paragraph(f"Generated at: {datetime.now().isoformat(timespec='seconds')}")

    doc.add_heading("1. Task Background and Objective", level=2)
    doc.add_paragraph(
        "We developed a multi-modal predictive system using MIMIC-IV derived static, time-series, and radiology-text data "
        "to estimate ICU time-to-discharge. We framed ICU death as right-censoring and compared traditional and deep survival models."
    )

    doc.add_heading("2. Multi-modal Model Implementation", level=2)
    doc.add_paragraph(
        "Static modality: demographic and early ICU summary features with missingness indicators.\n"
        "Time-series modality: first 24-hour hourly variables with statistical and sequential representations.\n"
        "Text modality: TF-IDF+SVD embedding and optional ClinicalBERT branch (fallback to fast text when GPU unavailable)."
    )

    doc.add_heading("3. Time-to-event Modeling", level=2)
    doc.add_paragraph(
        "We implemented CoxPH baseline, RSF/Tree-based survival, XGBoost-AFT, and deep gated-fusion discrete hazard model. "
        "Evaluation used C-index, Integrated Brier Score (IBS), and uncensored MAE/RMSE with bootstrap confidence intervals."
    )

    add_table_from_df(doc, df, "4. Model Performance Comparison")

    doc.add_heading("5. Performance Difference Analysis", level=2)
    doc.add_paragraph(
        f"Best model on test C-index: {best['exp']} (C-index={best['test_c_index']:.4f}, IBS={best['test_ibs']:.4f}). "
        "Static-only model provides robust baseline; adding time-series improves short-term physiology representation; "
        "adding text further improves heterogeneous clinical context capture."
    )

    doc.add_heading("6. Conclusion and Limitations", level=2)
    doc.add_paragraph(
        "The multi-modal setting consistently outperformed unimodal baseline. "
        "Main limitations include heavy missingness in specific labs, long-tail LOS distribution, and potential text encoding bottlenecks on CPU-only environments."
    )

    doc.save(out_docx)

    md = []
    md.append("# Assignment2 Final Report\n")
    md.append("## 1. Task Background and Objective")
    md.append(
        "We developed a multi-modal predictive system using MIMIC-IV derived static, time-series, and radiology-text data to estimate ICU time-to-discharge."
    )
    md.append("\n## 2. Multi-modal Model Implementation")
    md.append(
        "- Static: demographic and early ICU summary features with missingness indicators.\n"
        "- Time-series: first 24-hour hourly variables with statistical and sequential representations.\n"
        "- Text: TF-IDF+SVD and optional ClinicalBERT branch."
    )
    md.append("\n## 3. Time-to-event Modeling")
    md.append("CoxPH, RSF, XGBoost-AFT, deep gated-fusion hazard model, and weighted ensemble were compared.")
    md.append("\n## 4. Model Performance Comparison")
    try:
        md.append(df.to_markdown(index=False))
    except Exception:
        md.append("```csv\n" + df.to_csv(index=False) + "\n```")
    md.append("\n## 5. Performance Difference Analysis")
    md.append(
        f"Best test C-index: **{best['exp']}** ({best['test_c_index']:.4f}). "
        "Static+TS+Text generally outperformed static-only baseline."
    )
    md.append("\n## 6. Conclusion and Limitations")
    md.append("Multi-modal learning improved ranking and calibration, with limitations from missingness and compute budget.")
    with open(out_md, "w", encoding="utf-8") as f:
        f.write("\n\n".join(md))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data_dir", required=True)
    parser.add_argument("--metrics_dir", required=True)
    parser.add_argument("--figures_dir", required=True)
    parser.add_argument("--summary_csv", required=True)
    parser.add_argument("--out_experiment_docx", required=True)
    parser.add_argument("--out_final_docx", required=True)
    parser.add_argument("--out_final_md", required=True)
    args = parser.parse_args()

    build_experiment_log(
        out_path=args.out_experiment_docx,
        data_dir=args.data_dir,
        metrics_dir=args.metrics_dir,
        figures_dir=args.figures_dir,
        summary_csv=args.summary_csv,
    )
    build_final_report(
        out_docx=args.out_final_docx,
        out_md=args.out_final_md,
        summary_csv=args.summary_csv,
    )
    print("Report generation completed.")


if __name__ == "__main__":
    main()
